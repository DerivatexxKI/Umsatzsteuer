import streamlit as st
import pandas as pd
import io
from docx import Document
from io import BytesIO

st.set_page_config(page_title="UStVA KI-Assistent", layout="wide")
st.title("🤖 KI-Assistent für Umsatzsteuervoranmeldung")

# 1. Upload-Bereich
st.header("1️⃣ Buchungsdaten hochladen")
col1, col2 = st.columns(2)

with col1:
    uploaded_file = st.file_uploader("📁 Hauptbuchungsdaten (CSV/XLSX)", type=["csv", "xlsx"], key="main_upload")

with col2:
    uploaded_file2 = st.file_uploader("📄 Weitere Datei (z. B. Zusatzdaten, Mapping)", type=["csv", "xlsx"], key="extra_upload")

if uploaded_file:
    try:
        if uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file, sep=None, engine="python")
        else:
            df = pd.read_excel(uploaded_file)

        st.success("✅ Hauptdatei erfolgreich geladen")

        # Optional: zweite Datei anzeigen
        df2 = None
        if uploaded_file2:
            if uploaded_file2.name.endswith(".csv"):
                df2 = pd.read_csv(uploaded_file2, sep=None, engine="python")
            else:
                df2 = pd.read_excel(uploaded_file2)
            st.success("✅ Zusatzdatei erfolgreich geladen")
            with st.expander("📄 Zusatzdatei anzeigen"):
                st.dataframe(df2, use_container_width=True)

        # 2. USt-Kennziffern zuordnen
        st.header("2️⃣ USt-Kennziffern automatisch erkennen")
        def map_ustkennziffern(row):
            if row["Steuersatz"] == "19%":
                return "66"
            elif row["Steuersatz"] == "7%":
                return "81"
            elif row["Steuersatz"] == "steuerfrei":
                return "35"
            else:
                return "unbekannt"

        if "Steuersatz" in df.columns:
            df["USt-KZ"] = df.apply(map_ustkennziffern, axis=1)
        else:
            st.error("❌ Spalte 'Steuersatz' fehlt in der Datei")

        # 3. Visualisierung und Plausibilitätsprüfung
        st.header("3️⃣ Summen je USt-Kennziffer")
        grouped = df.groupby("USt-KZ")["Betrag"].sum().reset_index()
        grouped.columns = ["USt-KZ", "Summe EUR"]

        st.dataframe(grouped.style.format({"Summe EUR": "{:.2f}"}), use_container_width=True)
        st.bar_chart(grouped.set_index("USt-KZ"))

        with st.expander("🔍 Alle Buchungen im Detail"):
            st.dataframe(df, use_container_width=True)

        # 4. Abgleich mit Zusatzdatei
        merged = None
        if df2 is not None:
            st.header("4️⃣ Abgleich mit Zusatzdatei")
            common_cols = list(set(df.columns).intersection(df2.columns))
            if "Belegnummer" in common_cols:
                merged = df.merge(df2, on="Belegnummer", suffixes=("_Haupt", "_Zusatz"))
                st.success(f"🔗 {len(merged)} übereinstimmende Belegnummern gefunden")
                with st.expander("📎 Abgleich-Ergebnisse anzeigen"):
                    st.dataframe(merged, use_container_width=True)
            else:
                st.warning("⚠️ Kein gemeinsames Feld für Abgleich gefunden (z. B. 'Belegnummer')")

        # 5. Export als Word
        st.header("5️⃣ Export als Word-Datei")
        doc = Document()
        doc.add_heading("Umsatzsteuervoranmeldung – Zusammenfassung", 0)

        doc.add_heading("Summen je USt-Kennziffer", level=1)
        for _, row in grouped.iterrows():
            doc.add_paragraph(f"Kennziffer {row['USt-KZ']}: {row['Summe EUR']:.2f} EUR")

        if merged is not None:
            doc.add_heading("Abgleich mit Zusatzdatei", level=1)
            doc.add_paragraph(f"Übereinstimmende Belegnummern: {len(merged)}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Word-Dokument herunterladen",
            data=buffer,
            file_name="ustva_zusammenfassung.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"❌ Fehler beim Verarbeiten der Datei: {e}")
else:
    st.info("⬆️ Bitte lade eine Datei hoch, um zu starten")
